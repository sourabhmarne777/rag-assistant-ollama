import os
from typing import List, Dict, Any
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document types and split into chunks"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def process_document(self, file_path: str, source_name: str) -> List[Dict[str, Any]]:
        """Process a document file and return chunks"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            return self._split_text_into_chunks(text, source_name)
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def process_text(self, text: str, source_name: str) -> List[Dict[str, Any]]:
        """Process raw text and return chunks"""
        try:
            return self._split_text_into_chunks(text, source_name)
        except Exception as e:
            logger.error(f"Error processing text from {source_name}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if not text.strip():
                raise ValueError("The text file is empty")
            
            return text
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return text
            except Exception as e:
                logger.error(f"Error reading text file with latin-1 encoding: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error extracting text from TXT file: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _split_text_into_chunks(self, text: str, source_name: str) -> List[Dict[str, Any]]:
        """Split text into chunks using LangChain text splitter"""
        try:
            # Clean the text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create chunk objects with metadata
            chunk_objects = []
            for i, chunk in enumerate(chunks):
                if chunk.strip():  # Only add non-empty chunks
                    chunk_objects.append({
                        'content': chunk.strip(),
                        'chunk_id': f"{source_name}_chunk_{i}",
                        'source': source_name,
                        'chunk_index': i,
                        'total_chunks': len(chunks)
                    })
            
            logger.info(f"Split text from {source_name} into {len(chunk_objects)} chunks")
            return chunk_objects
            
        except Exception as e:
            logger.error(f"Error splitting text into chunks: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        try:
            # Remove excessive whitespace
            text = ' '.join(text.split())
            
            # Remove excessive newlines
            text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
            
            # Remove special characters that might cause issues
            text = text.replace('\x00', '')  # Remove null characters
            text = text.replace('\ufeff', '')  # Remove BOM
            
            return text
            
        except Exception as e:
            logger.error(f"Error cleaning text: {str(e)}")
            return text  # Return original text if cleaning fails
    
    def update_chunk_settings(self, chunk_size: int, chunk_overlap: int):
        """Update chunk size and overlap settings"""
        try:
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
            
            # Reinitialize text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
            
            logger.info(f"Updated chunk settings: size={chunk_size}, overlap={chunk_overlap}")
            
        except Exception as e:
            logger.error(f"Error updating chunk settings: {str(e)}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return ['.pdf', '.txt', '.docx']
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file can be processed"""
        try:
            if not os.path.exists(file_path):
                return False
            
            file_extension = os.path.splitext(file_path)[1].lower()
            return file_extension in self.get_supported_formats()
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {str(e)}")
            return False